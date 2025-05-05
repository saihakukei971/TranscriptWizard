# main.py - メインスクリプト
import os
import sys
import time
import csv
import whisper
import json
from datetime import datetime
from openai import OpenAI
from pydub import AudioSegment
from docx import Document
from dotenv import load_dotenv
from pathlib import Path
import subprocess
import logging
import traceback

# ロギング設定
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("transcriber.log"),
        logging.StreamHandler()
    ]
)

# スクリプトの初期化
def initialize():
    # 設定ファイルの読み込み
    if not os.path.exists(".env"):
        create_default_env()
    
    # .envファイルの読み込み
    load_dotenv()
    
    # APIキーの取得
    openai_key = os.getenv("OPENAI_API_KEY")
    slack_webhook = os.getenv("SLACK_WEBHOOK_URL")
    
    # APIキーが設定されていない場合は入力を求める
    if not openai_key or openai_key == "sk-XXXXXXXXXXXXXXXXXXXXXXXXXXXX":
        openai_key = input("OpenAI APIキーを入力してください: ")
        # .envファイルに保存
        with open(".env", "w", encoding="utf-8") as f:
            f.write(f"OPENAI_API_KEY={openai_key}\n")
            if slack_webhook:
                f.write(f"SLACK_WEBHOOK_URL={slack_webhook}\n")
    
    # FFmpegのチェック
    if not check_ffmpeg():
        print("警告: ffmpegが見つかりません。音声処理に必要です。")
        print("Windows: https://ffmpeg.org/download.html からダウンロードするか、")
        print("'choco install ffmpeg -y' でインストールしてください。")
        print("Mac: 'brew install ffmpeg' でインストールしてください。")
        print("\nffmpegのインストール後、再度実行してください。")
        input("Enterキーを押すと終了します...")
        sys.exit(1)
    
    # 処理ログの初期化
    init_log_file()
    
    return openai_key, slack_webhook

# デフォルトの.envファイルを作成
def create_default_env():
    with open(".env", "w", encoding="utf-8") as f:
        f.write("OPENAI_API_KEY=sk-XXXXXXXXXXXXXXXXXXXXXXXXXXXX\n")
        f.write("SLACK_WEBHOOK_URL=https://hooks.slack.com/services/XXX/YYY/ZZZ\n")
    logging.info(".envファイルを新規作成しました")

# ffmpegのチェック
def check_ffmpeg():
    try:
        subprocess.run(
            ["ffmpeg", "-version"], 
            stdout=subprocess.PIPE, 
            stderr=subprocess.PIPE, 
            check=True
        )
        return True
    except (subprocess.SubprocessError, FileNotFoundError):
        return False

# Whisperモデルのロード
def load_whisper_model(model_name="small"):
    try:
        print(f"Whisperモデル '{model_name}' をロード中...")
        logging.info(f"Whisperモデル '{model_name}' をロード中...")
        model = whisper.load_model(model_name)
        print(f"Whisperモデル '{model_name}' のロードが完了しました")
        logging.info(f"Whisperモデル '{model_name}' のロードが完了しました")
        return model
    except Exception as e:
        error_msg = f"Whisperモデルのロード中にエラーが発生しました: {e}"
        print(error_msg)
        logging.error(error_msg)
        return None

# 音声ファイルの処理
def process_audio_file(file_path, model, openai_key, slack_webhook, interval=5*60*1000):
    try:
        # 開始時間の記録
        start_time = datetime.now()
        basename = os.path.basename(file_path)
        print(f"\n{basename} の処理を開始します...")
        logging.info(f"ファイル '{file_path}' の処理を開始しました...")
        
        # 音声の読み込み
        audio = AudioSegment.from_file(file_path)
        
        # 5分ごとに分割
        segments = [audio[i:i+interval] for i in range(0, len(audio), interval)]
        
        # 文字起こし
        all_text = ""
        for i, segment in enumerate(segments):
            progress = f"セグメント {i+1}/{len(segments)}"
            print(f"{progress} を処理中...")
            logging.info(f"{progress} を処理中...")
            
            segment_file = f"chunk_{i}.mp3"
            segment.export(segment_file, format="mp3")
            
            result = model.transcribe(segment_file, language="Japanese")
            all_text += result["text"] + "\n\n"
            
            # 一時ファイルの削除
            os.remove(segment_file)
        
        # 文字起こし結果をファイルに保存
        output_file_path = f"{os.path.splitext(file_path)[0]}_文字起こし.txt"
        with open(output_file_path, "w", encoding="utf-8") as f:
            f.write(all_text)
        
        print(f"文字起こしが完了しました: {os.path.basename(output_file_path)}")
        logging.info(f"文字起こしが完了しました: {output_file_path}")
        
        # GPTによる要約
        print("GPTによる要約を実行中...")
        
        prompt = f"""
        以下の会議内容を、【議題】【要点】【決定事項】【ToDo】形式で簡潔にまとめてください。

        【会議内容】
        {all_text}
        """
        
        # OpenAI クライアントの初期化
        client = OpenAI(api_key=openai_key)
        
        try:
            # まずgpt-4.1を試し、失敗したらgpt-3.5-turboを使用
            try:
                response = client.chat.completions.create(
                    model="gpt-4o", 
                    messages=[{"role": "user", "content": prompt}]
                )
                summary_text = response.choices[0].message.content
                model_used = "gpt-4o"
            except:
                try:
                    response = client.chat.completions.create(
                        model="gpt-3.5-turbo", 
                        messages=[{"role": "user", "content": prompt}]
                    )
                    summary_text = response.choices[0].message.content
                    model_used = "gpt-3.5-turbo"
                except Exception as e:
                    logging.error(f"GPT APIリクエスト中にエラーが発生しました: {e}")
                    summary_text = "要約に失敗しました。文字起こしのみご利用ください。"
                    model_used = "失敗"
        except Exception as e:
            logging.error(f"GPTによる要約中にエラーが発生しました: {e}")
            summary_text = "要約に失敗しました。文字起こしのみご利用ください。"
            model_used = "失敗"
        
        # Word文書として保存
        docx_file_path = f"{os.path.splitext(file_path)[0]}_議事録.docx"
        doc = Document()
        doc.add_heading("議事録", 0)
        
        # 作成日時と元ファイル名を追加
        doc.add_paragraph(f"作成日時: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"元ファイル: {os.path.basename(file_path)}")
        doc.add_paragraph(f"使用モデル: Whisper({model.name})、要約({model_used})")
        doc.add_paragraph("")  # 空行
        
        # 要約内容を追加
        for line in summary_text.split("\n"):
            if line.strip() == "":
                doc.add_paragraph("")  # 空行を追加
            else:
                p = doc.add_paragraph()
                # 見出し（【】で囲まれた部分）を太字にする
                if "【" in line and "】" in line:
                    start_idx = line.find("【")
                    end_idx = line.find("】") + 1
                    if start_idx == 0:  # 行の先頭が【で始まる場合
                        p.add_run(line[:end_idx]).bold = True
                        if end_idx < len(line):
                            p.add_run(line[end_idx:])
                    else:
                        p.add_run(line)
                else:
                    p.add_run(line)
        
        doc.save(docx_file_path)
        print(f"議事録を保存しました: {os.path.basename(docx_file_path)}")
        logging.info(f"議事録を保存しました: {docx_file_path}")
        
        # 終了時間の記録
        end_time = datetime.now()
        processing_time = end_time - start_time
        
        # 処理ログの更新
        update_log(
            file_path=file_path,
            start_time=start_time,
            end_time=end_time,
            status="完了",
            text_length=len(all_text),
            summary_length=len(summary_text),
            slack_status="未送信"
        )
        
        # Slack通知
        if slack_webhook and slack_webhook != "https://hooks.slack.com/services/XXX/YYY/ZZZ":
            try:
                # 要約の先頭部分を抽出（最初の200字程度）
                summary_excerpt = summary_text[:200] + "..." if len(summary_text) > 200 else summary_text
                
                message = f"""
                *文字起こし・議事録作成が完了しました*
                ファイル: {os.path.basename(file_path)}
                処理時間: {processing_time.total_seconds() / 60:.1f}分
                使用モデル: Whisper({model.name})、要約({model_used})
                
                *議事録の抜粋:*
                ```
                {summary_excerpt}
                ```
                """
                
                success = send_to_slack(message, slack_webhook)
                
                # ログ更新
                if success:
                    update_log_field(file_path, "slack_status", "成功")
                    print("Slack通知を送信しました")
                    logging.info("Slack通知を送信しました")
                else:
                    update_log_field(file_path, "slack_status", "失敗")
                    print("Slack通知の送信に失敗しました")
                    logging.error("Slack通知の送信に失敗しました")
            except Exception as e:
                logging.error(f"Slack通知中にエラーが発生しました: {e}")
                update_log_field(file_path, "slack_status", "エラー")
        
        return True, processing_time.total_seconds() / 60
    
    except Exception as e:
        error_details = traceback.format_exc()
        error_msg = f"ファイル '{os.path.basename(file_path)}' の処理中にエラーが発生しました: {e}\n{error_details}"
        print(error_msg)
        logging.error(error_msg)
        
        end_time = datetime.now()
        
        # エラーログの更新
        update_log(
            file_path=file_path,
            start_time=start_time if 'start_time' in locals() else datetime.now(),
            end_time=end_time,
            status=f"エラー: {str(e)}",
            text_length=0,
            summary_length=0,
            slack_status="未送信"
        )
        
        return False, 0

# ログのCSVファイルを初期化（存在しない場合）
def init_log_file():
    log_file = "transcription_log.csv"
    if not os.path.exists(log_file):
        with open(log_file, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow([
                "ファイル名", "開始時刻", "終了時刻", "処理ステータス", 
                "文字数", "要約文字数", "Slack送信"
            ])
        logging.info("ログファイルを初期化しました")

# ログの更新
def update_log(file_path, start_time, end_time, status, text_length, summary_length, slack_status):
    try:
        with open("transcription_log.csv", "a", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow([
                os.path.basename(file_path),
                start_time.strftime("%Y-%m-%d %H:%M:%S"),
                end_time.strftime("%Y-%m-%d %H:%M:%S"),
                status,
                text_length,
                summary_length,
                slack_status
            ])
        
        logging.info(f"ログを更新しました: {os.path.basename(file_path)}")
    except Exception as e:
        logging.error(f"ログ更新中にエラーが発生しました: {e}")

# 特定のログフィールドの更新
def update_log_field(file_path, field, value):
    log_file = "transcription_log.csv"
    
    try:
        # CSVファイルの読み込み
        rows = []
        field_index = None
        
        with open(log_file, "r", encoding="utf-8-sig", newline="") as f:
            reader = csv.reader(f)
            headers = next(reader)
            
            # フィールドのインデックスを取得
            for i, header in enumerate(headers):
                if header.lower() == field.lower():
                    field_index = i
                    break
            
            if field_index is None:
                logging.error(f"フィールド '{field}' がログファイルに見つかりません")
                return False
            
            # 行を読み込む
            for row in reader:
                if row[0] == os.path.basename(file_path):
                    row[field_index] = value
                rows.append(row)
        
        # CSVファイルに書き戻す
        with open(log_file, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            writer.writerows(rows)
        
        return True
    
    except Exception as e:
        logging.error(f"ログフィールド更新中にエラーが発生しました: {e}")
        return False

# Slackにメッセージを送信
def send_to_slack(message, webhook_url):
    try:
        import requests
        
        payload = {"text": message}
        headers = {"Content-Type": "application/json"}
        response = requests.post(webhook_url, data=json.dumps(payload), headers=headers)
        
        if response.status_code == 200:
            return True
        else:
            logging.error(f"Slack通知の送信に失敗しました。ステータスコード: {response.status_code}")
            return False
    
    except Exception as e:
        logging.error(f"Slack通知の送信中にエラーが発生しました: {e}")
        return False

# メイン処理
def main():
    print("音声文字起こし・議事録作成ツール v1.0.0")
    print("=====================================")
    logging.info("音声文字起こし・議事録作成ツールを起動しました")
    
    try:
        # 初期化
        openai_key, slack_webhook = initialize()
        
        # Whisperモデルのロード
        model = load_whisper_model("small")
        if model is None:
            print("Whisperモデルのロードに失敗しました。プログラムを終了します。")
            input("Enterキーを押すと終了します...")
            sys.exit(1)
        
        # オーディオファイルの検索
        audio_extensions = [".mp3", ".wav", ".m4a", ".mp4"]
        audio_files = []
        
        for file in sorted(os.listdir(".")):
            if any(file.lower().endswith(ext) for ext in audio_extensions):
                audio_files.append(file)
        
        if not audio_files:
            print("\n処理対象の音声ファイルが見つかりません。")
            print("mp3, wav, m4a, mp4形式のファイルをこのフォルダに配置してから再度実行してください。")
            input("\nEnterキーを押すと終了します...")
            sys.exit(0)
        
        # ファイル数の表示
        file_count = len(audio_files)
        print(f"\n{file_count}個の音声ファイルが見つかりました:")
        for i, file in enumerate(audio_files, 1):
            print(f"{i}. {file}")
        
        print("\n処理を開始します...")
        logging.info(f"{file_count}個の音声ファイルの処理を開始します")
        
        # 各ファイルの処理
        successful_files = 0
        total_processing_time = 0
        
        for i, file in enumerate(audio_files, 1):
            print(f"\n[{i}/{file_count}] {file} を処理中...")
            success, processing_time = process_audio_file(file, model, openai_key, slack_webhook)
            
            if success:
                successful_files += 1
                total_processing_time += processing_time
                print(f"{file} の処理が完了しました。処理時間: {processing_time:.1f}分")
            else:
                print(f"{file} の処理中にエラーが発生しました。詳細はログを確認してください。")
        
        # 処理完了メッセージ
        print("\n======== 処理完了 ========")
        print(f"処理結果: {successful_files}/{file_count} ファイルの処理に成功しました")
        if successful_files > 0:
            print(f"合計処理時間: {total_processing_time:.1f}分")
        
        print("\n出力ファイル:")
        for file in audio_files:
            base_name = os.path.splitext(file)[0]
            txt_file = f"{base_name}_文字起こし.txt"
            docx_file = f"{base_name}_議事録.docx"
            
            if os.path.exists(txt_file):
                print(f"- {txt_file}")
            if os.path.exists(docx_file):
                print(f"- {docx_file}")
        
        print("\n処理ログ: transcription_log.csv")
        
        logging.info("すべての処理が完了しました")
    
    except Exception as e:
        error_details = traceback.format_exc()
        error_msg = f"予期せぬエラーが発生しました: {e}\n{error_details}"
        print(f"\nエラー: {error_msg}")
        logging.error(error_msg)
    
    finally:
        print("\n処理を終了します。")
        input("Enterキーを押すと終了します...")

if __name__ == "__main__":
    main()

# setup.py - exe作成に必要な情報を提供するスクリプト
"""
PyInstallerを使用して実行ファイルを作成するためのスクリプト

使用方法:
1. pip install pyinstaller で PyInstaller をインストール
2. pyinstaller setup.py を実行

これにより、dist/Transcriber.exe が生成されます。
"""

from PyInstaller.__main__ import run

if __name__ == "__main__":
    opts = [
        '--name=Transcriber',
        '--onefile',
        '--console',
        '--add-data=ffmpeg;ffmpeg',  # ffmpegを同梱
        '--clean',
        'main.py'
    ]
    run(opts)

# requirements.txt - 必要なライブラリ一覧
"""
openai==1.11.0
git+https://github.com/openai/whisper.git
pydub==0.25.1
python-docx==0.8.11
python-dotenv==1.0.0
requests==2.31.0
pyinstaller==6.0.0
"""

# .env.template - 環境変数テンプレート
"""
OPENAI_API_KEY=sk-XXXXXXXXXXXXXXXXXXXXXXXXXXXX
SLACK_WEBHOOK_URL=https://hooks.slack.com/services/XXX/YYY/ZZZ
"""